from __future__ import annotations
import logging
import json
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)

@dataclass
class ResumeHeader:
    name: str = ""
    headline: str = ""
    phone: str = ""
    email: str = ""
    portfolio_url: str = ""
    linkedin_url: str = ""

@dataclass
class CompetencyGroup:
    label: str
    items: List[str]

@dataclass
class ExperienceItem:
    role: str
    company: str
    location: str
    dates: str
    bullets: List[str]

@dataclass
class ProjectItem:
    title: str
    description: str
    link: Optional[str] = None

@dataclass
class EducationItem:
    degree: str
    school: str
    year: str

@dataclass
class ResumeContent:
    header: ResumeHeader = field(default_factory=ResumeHeader)
    summary: str = ""
    core_competencies: List[CompetencyGroup] = field(default_factory=list)
    experience: List[ExperienceItem] = field(default_factory=list)
    projects: List[ProjectItem] = field(default_factory=list)
    education: List[EducationItem] = field(default_factory=list)
    awards: List[str] = field(default_factory=list)
    template_name: str = "Andy Warthog"
    last_updated: str = field(default_factory=lambda: "")

@dataclass
class ValidationIssue:
    severity: str  # 'info', 'warning', 'critical'
    message: str

class AndyWarthogRenderer:
    """Renders structured JSON resume data into the Andy Warthog style (HTML/CSS)."""

    STYLE_CSS = """
    @import url('https://fonts.googleapis.com/css2?family=Open+Sans:wght@300;400;600;700&display=swap');
    
    body {
        font-family: 'Open Sans', sans-serif;
        color: #333;
        line-height: 1.4;
        margin: 0;
        padding: 40px;
        background: white;
    }
    .container {
        max-width: 800px;
        margin: auto;
    }
    header {
        margin-bottom: 20px;
    }
    h1 {
        margin: 0;
        font-size: 36px;
        color: #00695c; /* Teal */
        font-weight: 700;
        text-transform: none;
    }
    .headline {
        font-size: 18px;
        color: #00897b;
        font-style: italic;
        margin-bottom: 10px;
    }
    .contact-info {
        display: flex;
        justify-content: flex-end;
        gap: 15px;
        font-size: 12px;
        margin-top: -30px;
    }
    .contact-info a {
        color: #00695c;
        text-decoration: none;
    }
    
    h2 {
        font-size: 18px;
        color: #00695c;
        text-transform: uppercase;
        border-bottom: 1px solid #b2dfdb;
        padding-bottom: 5px;
        margin-top: 25px;
        margin-bottom: 10px;
    }
    
    .summary {
        font-size: 12px;
        margin-bottom: 20px;
    }
    
    .competencies {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 20px;
        margin-bottom: 20px;
    }
    .comp-group h3 {
        font-size: 13px;
        color: #00897b;
        margin-bottom: 5px;
    }
    .comp-group p {
        font-size: 11px;
        margin: 0;
    }
    
    .exp-item {
        margin-bottom: 15px;
    }
    .exp-header {
        display: flex;
        justify-content: space-between;
        align-items: baseline;
    }
    .exp-title {
        font-weight: 700;
        font-size: 13px;
        color: #00897b;
    }
    .exp-company {
        font-size: 12px;
        font-style: italic;
    }
    .exp-dates {
        font-size: 11px;
        color: #666;
    }
    ul {
        margin: 5px 0;
        padding-left: 20px;
    }
    li {
        font-size: 11px;
        margin-bottom: 3px;
    }
    
    .projects {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 20px;
    }
    .project-item h3 {
        font-size: 13px;
        color: #00897b;
        margin-bottom: 5px;
    }
    .project-item p {
        font-size: 11px;
        margin-bottom: 5px;
    }
    .project-link {
        font-size: 10px;
        color: #00695c;
        text-decoration: none;
        font-weight: 600;
    }
    
    .edu-item {
        margin-bottom: 10px;
    }
    .edu-degree {
        font-weight: 700;
        font-size: 12px;
    }
    .edu-school {
        font-size: 11px;
    }
    
    .awards-list {
        display: grid;
        grid-template-columns: 1fr;
        gap: 5px;
    }
    
    @media print {
        body { padding: 0; }
        .container { max-width: 100%; }
    }
    """

    def normalize(self, content: ResumeContent) -> tuple[ResumeContent, List[ValidationIssue]]:
        """Apply normalization rules and return (normalized_content, issues)."""
        issues = []
        
        # 1. Structural Normalization
        if len(content.summary) > 450:
            content.summary = content.summary[:447] + "..."
            issues.append(ValidationIssue('info', "Summary was truncated to 450 characters to fit the template."))
        elif len(content.summary) < 50:
            issues.append(ValidationIssue('warning', "Summary is very short; consider elaborating."))
            
        if not content.experience:
            issues.append(ValidationIssue('critical', "No professional experience entries found."))
        
        for item in content.experience:
            if len(item.bullets) > 3:
                item.bullets = item.bullets[:3]
                issues.append(ValidationIssue('info', f"Role '{item.role}' has >3 bullets; only the first 3 are shown."))
            
            for j, bullet in enumerate(item.bullets):
                if len(bullet) > 163:
                    item.bullets[j] = bullet[:160] + "..."
                    issues.append(ValidationIssue('info', f"A bullet in '{item.company}' was truncated."))
            
        if len(content.projects) > 2:
            content.projects = content.projects[:2]
            issues.append(ValidationIssue('warning', "Only 2 featured projects are supported; others are hidden."))
            
        if not content.header.name:
            issues.append(ValidationIssue('critical', "Name is missing in the header."))
        if not content.header.email and not content.header.phone:
            issues.append(ValidationIssue('critical', "At least one contact method (email or phone) is required."))

        # 2. Content Quality Validation (Placeholders, Tokens, Filler)
        content_issues = self._validate_content_quality(content)
        issues.extend(content_issues)
            
        return content, issues

    def _validate_content_quality(self, content: ResumeContent) -> List[ValidationIssue]:
        issues = []
        import re
        
        # Patterns
        placeholders = r"\[.*?\]|TODO|TBD|INSERT|LOREM IPSUM"
        tokens = r"\{.*?\}|\{\{.*?\}\}"
        generic_filler = [
            "detail-oriented professional",
            "results-driven leader",
            "excellent communication skills",
            "team player"
        ]

        # Check all string fields in header
        h = content.header
        for field_val in [h.name, h.headline, h.email, h.phone]:
            if re.search(placeholders, field_val, re.I):
                issues.append(ValidationIssue('critical', f"Header contains placeholder: '{field_val}'"))
            if re.search(tokens, field_val):
                issues.append(ValidationIssue('critical', f"Header contains unresolved token: '{field_val}'"))

        # Check summary
        if re.search(placeholders, content.summary, re.I):
            issues.append(ValidationIssue('critical', "Summary contains placeholder text."))
        if re.search(tokens, content.summary):
            issues.append(ValidationIssue('critical', "Summary contains unresolved tokens."))
        
        filler_count = sum(1 for f in generic_filler if f in content.summary.lower())
        if filler_count >= 2:
            issues.append(ValidationIssue('warning', "Summary contains multiple generic buzzwords; consider more specific achievements."))

        # Check experience bullets
        for exp in content.experience:
            for bullet in exp.bullets:
                if re.search(placeholders, bullet, re.I):
                    issues.append(ValidationIssue('critical', f"Experience ({exp.company}) contains placeholder."))
                if re.search(tokens, bullet):
                    issues.append(ValidationIssue('critical', f"Experience ({exp.company}) contains unresolved token."))
                
        return issues

    def to_html(self, content: ResumeContent, ats_safe: bool = False) -> str:
        """Render content to HTML string."""
        content, _ = self.normalize(content)
        
        # Simple template rendering
        comp_html = "".join([
            f'<div class="comp-group"><h3>{g.label}</h3><p>{", ".join(g.items)}</p></div>'
            for g in content.core_competencies
        ])
        
        exp_html = "".join([
            f'''
            <div class="exp-item">
                <div class="exp-header">
                    <span class="exp-title">{item.role}</span>
                    <span class="exp-dates">{item.dates}</span>
                </div>
                <div class="exp-company">{item.company} | {item.location}</div>
                <ul>{"".join([f"<li>{b}</li>" for b in item.bullets])}</ul>
            </div>
            ''' for item in content.experience
        ])
        
        proj_html = "".join([
            f'''
            <div class="project-item">
                <h3>{p.title}</h3>
                <p>{p.description}</p>
                {f'<a href="{p.link}" class="project-link">View Project</a>' if p.link else ''}
            </div>
            ''' for p in content.projects
        ])
        
        edu_html = "".join([
            f'''
            <div class="edu-item">
                <div class="edu-degree">{e.degree}</div>
                <div class="edu-school">{e.school}, {e.year}</div>
            </div>
            ''' for e in content.education
        ])
        
        awards_html = "".join([f"<li>{a}</li>" for a in content.awards])
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                {self.STYLE_CSS if not ats_safe else "body { font-family: sans-serif; padding: 20px; } h2 { border-bottom: 1px solid #ccc; }"}
            </style>
        </head>
        <body>
            <div class="container">
                <header>
                    <h1>{content.header.name}</h1>
                    <div class="headline">{content.header.headline}</div>
                    <div class="contact-info">
                        <span>{content.header.phone}</span>
                        <a href="mailto:{content.header.email}">{content.header.email}</a>
                        {f'<a href="{content.header.portfolio_url}">Portfolio</a>' if content.header.portfolio_url else ''}
                        {f'<a href="{content.header.linkedin_url}">LinkedIn</a>' if content.header.linkedin_url else ''}
                    </div>
                </header>
                
                <section>
                    <h2>Professional Summary</h2>
                    <div class="summary">{content.summary}</div>
                </section>
                
                <section>
                    <h2>Core Competencies</h2>
                    <div class="competencies">
                        {comp_html}
                    </div>
                </section>
                
                <section>
                    <h2>Professional Experience</h2>
                    {exp_html}
                </section>
                
                <section>
                    <h2>Featured Projects</h2>
                    <div class="projects">
                        {proj_html}
                    </div>
                </section>
                
                <section>
                    <h2>Education</h2>
                    {edu_html}
                </section>
                
                {f'<section><h2>Awards & Recognition</h2><ul>{awards_html}</ul></section>' if content.awards else ''}
            </div>
        </body>
        </html>
        """
        return html

    async def render_pdf(self, html: str, output_path: Path):
        """Render HTML to PDF using Playwright."""
        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.set_content(html)
            await page.pdf(path=str(output_path), format="A4", print_background=True)
            await browser.close()

    def render_docx(self, content: ResumeContent, output_path: Path):
        """Render ResumeContent to a DOCX file using python-docx."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        content, _ = self.normalize(content)
        doc = Document()
        
        # Helper for teal color
        TEAL = RGBColor(0, 105, 92)
        LIGHT_TEAL = RGBColor(0, 137, 123)

        # Header
        h = doc.add_heading(content.header.name, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = TEAL
            
        p_headline = doc.add_paragraph()
        run_headline = p_headline.add_run(content.header.headline)
        run_headline.italic = True
        run_headline.font.size = Pt(14)
        run_headline.font.color.rgb = LIGHT_TEAL
        
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        contact_parts = [content.header.phone, content.header.email]
        if content.header.portfolio_url: contact_parts.append(content.header.portfolio_url)
        if content.header.linkedin_url: contact_parts.append(content.header.linkedin_url)
        p_contact.add_run(" | ".join(filter(None, contact_parts))).font.size = Pt(9)

        # Sections
        def add_section_header(text):
            hdr = doc.add_heading(text, level=1)
            for run in hdr.runs:
                run.font.color.rgb = TEAL
                run.font.size = Pt(14)

        # Summary
        add_section_header("Professional Summary")
        doc.add_paragraph(content.summary).style.font.size = Pt(10)

        # Competencies
        add_section_header("Core Competencies")
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        row_cells = table.rows[0].cells
        for i, group in enumerate(content.core_competencies[:3]):
            cell_p = row_cells[i].paragraphs[0]
            r_label = cell_p.add_run(f"{group.label}:\n")
            r_label.bold = True
            r_label.font.color.rgb = LIGHT_TEAL
            cell_p.add_run(", ".join(group.items)).font.size = Pt(9)

        # Experience
        add_section_header("Professional Experience")
        for item in content.experience:
            p_exp = doc.add_paragraph()
            r_role = p_exp.add_run(item.role)
            r_role.bold = True
            r_role.font.color.rgb = LIGHT_TEAL
            p_exp.add_run(f"\t{item.dates}").font.size = Pt(10)
            p_exp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p_co = doc.add_paragraph()
            r_co = p_co.add_run(f"{item.company} | {item.location}")
            r_co.italic = True
            r_co.font.size = Pt(10)
            
            for bullet in item.bullets:
                doc.add_paragraph(bullet, style='List Bullet').style.font.size = Pt(9)

        # Projects
        if content.projects:
            add_section_header("Featured Projects")
            for proj in content.projects:
                p_proj = doc.add_paragraph()
                r_title = p_proj.add_run(proj.title)
                r_title.bold = True
                r_title.font.color.rgb = LIGHT_TEAL
                doc.add_paragraph(proj.description).style.font.size = Pt(9)
                if proj.link:
                    doc.add_paragraph(f"Link: {proj.link}").style.font.size = Pt(8)

        # Education
        add_section_header("Education")
        for edu in content.education:
            p_edu = doc.add_paragraph()
            p_edu.add_run(edu.degree).bold = True
            doc.add_paragraph(f"{edu.school}, {edu.year}").style.font.size = Pt(10)

        # Awards
        if content.awards:
            add_section_header("Awards & Recognition")
            for award in content.awards:
                doc.add_paragraph(award, style='List Bullet').style.font.size = Pt(9)

        doc.save(str(output_path))
